import contextlib
import os

import pytest
import typer
from docx import Document

from rendercv.cli.new_command.new_command import cli_command_new
from rendercv.schema.rendercv_model_builder import build_rendercv_dictionary_and_model


class TestCliCommandNew:
    @pytest.mark.parametrize("create_typst_templates", [True, False])
    @pytest.mark.parametrize("create_markdown_templates", [True, False])
    @pytest.mark.parametrize("input_file_exists", [True, False])
    @pytest.mark.parametrize("typst_templates_exist", [True, False])
    @pytest.mark.parametrize("markdown_templates_exist", [True, False])
    def test_cli_command_new(
        self,
        tmp_path,
        create_typst_templates,
        create_markdown_templates,
        input_file_exists,
        typst_templates_exist,
        markdown_templates_exist,
    ):
        os.chdir(tmp_path)

        full_name = "John Doe"
        input_file_path = tmp_path / "John_Doe_CV.yaml"
        typst_folder = tmp_path / "classic"
        markdown_folder = tmp_path / "markdown"

        # Set up pre-existing files/folders
        if input_file_exists:
            input_file_path.write_text("existing content", encoding="utf-8")
        if typst_templates_exist:
            typst_folder.mkdir()
        if markdown_templates_exist:
            markdown_folder.mkdir()

        cli_command_new(
            full_name=full_name,
            create_typst_templates=create_typst_templates,
            create_markdown_templates=create_markdown_templates,
        )

        # Input file should always exist after command
        assert input_file_path.exists()

        # If input file existed before, content should be unchanged
        if input_file_exists:
            assert input_file_path.read_text(encoding="utf-8") == "existing content"
        else:
            # Make sure it's a valid YAML file
            build_rendercv_dictionary_and_model(input_file_path)

        # Typst templates
        if not create_typst_templates:
            # Should not be created (unless already existed)
            if not typst_templates_exist:
                assert not typst_folder.exists()
        else:
            # Should exist (created or already existed)
            assert typst_folder.exists()
            if typst_templates_exist:
                assert not (typst_folder / "Header.j2.typ").exists()
            else:
                assert (typst_folder / "Header.j2.typ").exists()

        # Markdown templates
        if not create_markdown_templates:
            # Should not be created (unless already existed)
            if not markdown_templates_exist:
                assert not markdown_folder.exists()
        else:
            # Should exist (created or already existed)
            assert markdown_folder.exists()
            if markdown_templates_exist:
                assert not (markdown_folder / "Header.j2.md").exists()
            else:
                assert (markdown_folder / "Header.j2.md").exists()

    def test_errors_for_invalid_theme(self, capsys):
        with contextlib.suppress(typer.Exit):
            cli_command_new(
                full_name="John Doe",
                theme="invalid_theme",
                create_typst_templates=False,
                create_markdown_templates=False,
            )

            captured = capsys.readouterr()
            assert "Available themes are:" in captured.out

    def test_errors_for_invalid_locale(self, capsys):
        with contextlib.suppress(typer.Exit):
            cli_command_new(
                full_name="John Doe",
                locale="invalid_locale",
                create_typst_templates=False,
                create_markdown_templates=False,
            )

            captured = capsys.readouterr()
            assert "Available locales are:" in captured.out

    def test_accepts_docx_file(self, tmp_path):
        """Test that .docx files are accepted for CV extraction."""
        os.chdir(tmp_path)
        
        # Create a test .docx file with random data
        doc_path = tmp_path / "test_resume.docx"
        doc = Document()
        doc.add_paragraph("Alice Smith")
        doc.add_paragraph("Data Scientist")
        doc.add_paragraph("New York, NY")
        doc.add_paragraph("Email: alice.smith@example.com")
        doc.add_paragraph("Phone: +1-555-123-4567")
        doc.save(str(doc_path))
        
        # Should not raise an error
        cli_command_new(
            full_name="Alice Smith",
            word_file_path=str(doc_path),
            create_typst_templates=False,
            create_markdown_templates=False,
        )
        
        # Verify the YAML file was created
        input_file_path = tmp_path / "Alice_Smith_CV.yaml"
        assert input_file_path.exists()

    def test_rejects_txt_file(self, tmp_path, capsys):
        """Test that .txt files are rejected."""
        os.chdir(tmp_path)
        
        # Create a .txt file
        txt_path = tmp_path / "resume.txt"
        txt_path.write_text("This is a text file with CV data")
        
        with contextlib.suppress(typer.Exit):
            cli_command_new(
                full_name="Bob Johnson",
                word_file_path=str(txt_path),
                create_typst_templates=False,
                create_markdown_templates=False,
            )
        
        captured = capsys.readouterr()
        assert "Unsupported file format" in captured.out
        assert "Only .docx files" in captured.out

    def test_rejects_doc_file(self, tmp_path, capsys):
        """Test that .doc files are rejected."""
        os.chdir(tmp_path)
        
        # Create a .doc file
        doc_path = tmp_path / "resume.doc"
        doc_path.write_text("This is an old .doc file")
        
        with contextlib.suppress(typer.Exit):
            cli_command_new(
                full_name="Charlie Brown",
                word_file_path=str(doc_path),
                create_typst_templates=False,
                create_markdown_templates=False,
            )
        
        captured = capsys.readouterr()
        assert "Unsupported file format" in captured.out
        assert "Only .docx files" in captured.out

    def test_rejects_pdf_file(self, tmp_path, capsys):
        """Test that .pdf files are rejected."""
        os.chdir(tmp_path)
        
        # Create a .pdf file (just a text file with .pdf extension)
        pdf_path = tmp_path / "resume.pdf"
        pdf_path.write_bytes(b"PDF content would be here")
        
        with contextlib.suppress(typer.Exit):
            cli_command_new(
                full_name="Diana Prince",
                word_file_path=str(pdf_path),
                create_typst_templates=False,
                create_markdown_templates=False,
            )
        
        captured = capsys.readouterr()
        assert "Unsupported file format" in captured.out
        assert "Only .docx files" in captured.out

    def test_rejects_file_without_extension(self, tmp_path, capsys):
        """Test that files without extensions are rejected."""
        os.chdir(tmp_path)
        
        # Create a file without extension
        no_ext_path = tmp_path / "resume"
        no_ext_path.write_text("File without extension")
        
        with contextlib.suppress(typer.Exit):
            cli_command_new(
                full_name="Eve Wilson",
                word_file_path=str(no_ext_path),
                create_typst_templates=False,
                create_markdown_templates=False,
            )
        
        captured = capsys.readouterr()
        assert "Unsupported file format" in captured.out
        assert "Only .docx files" in captured.out

    def test_rejects_uppercase_docx_extension(self, tmp_path):
        """Test that .DOCX (uppercase) files are accepted (case insensitive)."""
        os.chdir(tmp_path)
        
        # Create a test .DOCX file with random data
        doc_path = tmp_path / "test_resume.DOCX"
        doc = Document()
        doc.add_paragraph("Frank Miller")
        doc.add_paragraph("Software Architect")
        doc.add_paragraph("Seattle, WA")
        doc.add_paragraph("Email: frank.miller@example.com")
        doc.save(str(doc_path))
        
        # Should not raise an error (case insensitive)
        cli_command_new(
            full_name="Frank Miller",
            word_file_path=str(doc_path),
            create_typst_templates=False,
            create_markdown_templates=False,
        )
        
        # Verify the YAML file was created
        input_file_path = tmp_path / "Frank_Miller_CV.yaml"
        assert input_file_path.exists()
